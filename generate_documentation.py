"""
Generate comprehensive documentation for the Testing Folder project
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_documentation():
    doc = Document()
    
    # Title
    title = doc.add_heading('Testing Folder - Complete Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project Overview
    doc.add_heading('🎯 Project Overview', 1)
    doc.add_paragraph(
        'This is an AI-Powered Sign Language Recognition and Translation System that combines:'
    )
    bullets = [
        '🤟 Real-time sign language recognition (MediaPipe + TensorFlow)',
        '🧠 AI text processing (Ollama Qwen2.5:7b)',
        '🎥 Sign language video generation and playback',
        '🗣️ Text-to-Speech synthesis (TTS)'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    # Folder Structure
    doc.add_heading('📂 Folder Structure', 1)
    
    # Core Files Section
    doc.add_heading('1. Core Program Files', 2)
    
    doc.add_heading('📄 app.py (916 lines) - Main Application', 3)
    doc.add_paragraph('Function: PyQt6 graphical interface + real-time sign language recognition')
    doc.add_paragraph('Key Components:', style='Heading 5')
    
    components = [
        ('CameraWorker (Lines 184-364)', [
            'Uses MediaPipe Hands to detect 21 hand keypoints',
            'KeyPointClassifier identifies A-Z letters',
            '0.6 second hold time confirms gestures',
            'Cursor-based text editing system'
        ]),
        ('PredictionWorker (Lines 58-130)', [
            'Ollama AI smart word prediction',
            'Auto-completes current word or predicts next word'
        ]),
        ('TTSWorker (Lines 136-177)', [
            'Text-to-Speech (pyttsx3)',
            'Background thread prevents UI freezing'
        ]),
        ('UnifiedApp (Lines 383-840)', [
            'Dual-tab interface:',
            '  - 🗣️ Tab 1: Live sign language communication',
            '  - 🎥 Tab 2: Text to sign language video'
        ])
    ]
    
    for comp_name, details in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(comp_name).bold = True
        for detail in details:
            doc.add_paragraph(detail, style='List Bullet 2')
    
    doc.add_paragraph('Technical Highlights:', style='Heading 5')
    highlights = [
        'PyQt signal-slot mechanism for thread-safe communication',
        'OpenCV video stitching (15-frame crossfade)',
        'AI grammar correction',
        'Semantic similarity word search'
    ]
    for highlight in highlights:
        doc.add_paragraph(highlight, style='List Bullet')
    
    doc.add_heading('📄 logic.py (150 lines) - AI Logic Module', 3)
    doc.add_paragraph('Functions:')
    logic_functions = [
        ('Speech Recognition (Lines 26-42)', 'listen_to_audio() → Microphone input to text'),
        ('ASL Gloss Translation (Lines 47-75)', 'translate_to_gloss("I want to eat") → ["I", "WANT", "EAT"]'),
        ('Video Path Lookup (Lines 80-112)', 'Supports variant filenames (hello_1.mp4, hello_2.mp4)'),
        ('Video Stitching (Lines 114-150)', 'MoviePy stitching with 0.2s crossfade')
    ]
    for func_name, desc in logic_functions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(func_name + ': ').bold = True
        p.add_run(desc)
    
    # AI Model Folder
    doc.add_heading('2. AI Model Folder', 2)
    doc.add_heading('📁 model/keypoint_classifier/', 3)
    model_files = [
        'keypoint_classifier.tflite    ← TensorFlow Lite gesture classification model',
        'keypoint_classifier.keras     ← Keras training version',
        'keypoint_classifier.py        ← Model loader',
        'keypoint_classifier_label.csv ← A-Z label mapping',
        'keypoint.csv                  ← Training data (21 keypoints × 2 coords = 42 dimensions)'
    ]
    for mf in model_files:
        doc.add_paragraph(mf, style='List Bullet')
    
    doc.add_paragraph('Working Principle:', style='Heading 5')
    principles = [
        'MediaPipe extracts 21 hand keypoints',
        'Normalized to wrist-relative coordinates',
        'TFLite model classifies → 0-25 index → A-Z letters'
    ]
    for principle in principles:
        doc.add_paragraph(f'{principles.index(principle) + 1}. {principle}')
    
    # Dataset Folders
    doc.add_heading('3. Dataset Folders', 2)
    doc.add_heading('📁 custom/ (2000 MP4 files)', 3)
    doc.add_paragraph('Your custom sign language video database organized by first letter:')
    doc.add_paragraph('custom/a/ → about.mp4, after.mp4, ...', style='List Bullet')
    doc.add_paragraph('custom/b/ → baby.mp4, bad.mp4, ...', style='List Bullet')
    doc.add_paragraph('custom/z/ → zoo.mp4', style='List Bullet')
    
    doc.add_heading('📁 dataset/wlasl/', 3)
    doc.add_paragraph('WLASL (Word-Level American Sign Language) dataset containing common sign language videos')
    
    doc.add_heading('📄 WLASL_v0.3.json', 3)
    doc.add_paragraph('Dataset index file with video metadata and frame information')
    
    # Utility Folders
    doc.add_heading('4. Utility Folders', 2)
    
    doc.add_heading('📁 file_processing/', 3)
    processing_files = [
        ('setup_dataset.py', 'Organizes videos from WLASL JSON into folder structure'),
        ('select_best_videos.py', 'Selects best video from multiple variants')
    ]
    for pf_name, pf_desc in processing_files:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(pf_name + ': ').bold = True
        p.add_run(pf_desc)
    
    doc.add_heading('📁 utils/', 3)
    doc.add_paragraph('cvfpscalc.py: FPS calculation tool for performance monitoring', style='List Bullet')
    
    doc.add_heading('📁 avatars/', 3)
    avatar_files = [
        'avatar_generator_astroboy.py  ← Astro Boy virtual avatar generator',
        'avatar_generator_ironman      ← Iron Man virtual avatar',
        'face_generator_ironman_body   ← Face swap + Iron Man body',
        'inswapper_128.onnx            ← InsightFace face swap model',
        'real_face_expression          ← Real expression generation'
    ]
    for af in avatar_files:
        doc.add_paragraph(af, style='List Bullet')
    doc.add_paragraph('Purpose: Swap real person sign language videos with virtual characters (Iron Man, Astro Boy)')
    
    # Output Folder
    doc.add_heading('5. Output Folder', 2)
    doc.add_heading('📁 output/', 3)
    doc.add_paragraph('Stores generated sign language translation videos with timestamp filenames to avoid conflicts')
    doc.add_paragraph('Example: output_1702835621.mp4', style='List Bullet')
    
    # Configuration
    doc.add_heading('6. Configuration & Dependencies', 2)
    doc.add_heading('📄 requirement.txt', 3)
    doc.add_paragraph('Project dependencies:')
    requirements = [
        'PyQt6         ← GUI framework',
        'opencv-python ← Image processing',
        'mediapipe     ← Hand detection',
        'tensorflow    ← Model inference',
        'requests      ← Ollama API communication',
        'pyttsx3       ← Text-to-Speech',
        'pyqtdarktheme ← Dark theme',
        'moviepy       ← Video editing'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    # Workflows
    doc.add_page_break()
    doc.add_heading('🔄 Complete Workflows', 1)
    
    doc.add_heading('Scenario 1: Real-time Sign Recognition → Text → Speech', 2)
    workflow1 = [
        '1. Camera captures hand gesture',
        '2. MediaPipe detects hand keypoints (21 points)',
        '3. KeyPointClassifier identifies letter "H"',
        '4. Hold for 0.6 seconds → Confirmed',
        '5. Insert character to text buffer "Hello"',
        '6. Ollama AI predicts next word ["world", "there", "everyone"]',
        '7. User presses F1 to select "world"',
        '8. Press Enter → AI corrects grammar "Hello world!"',
        '9. TTS speaks output 🔊'
    ]
    for step in workflow1:
        doc.add_paragraph(step)
    
    doc.add_heading('Scenario 2: Text → Sign Language Video', 2)
    workflow2 = [
        '1. User inputs "I want to eat lunch"',
        '2. Ollama AI converts to ASL Gloss: ["I", "WANT", "EAT", "LUNCH"]',
        '3. Search for videos:',
        '   - custom/i/i.mp4 ✓',
        '   - custom/w/want.mp4 ✓',
        '   - custom/e/eat.mp4 ✓',
        '   - custom/l/lunch.mp4 ✗ (missing)',
        '4. AI finds similar word: "lunch" → "food" ✓',
        '5. OpenCV stitches videos (15-frame crossfade)',
        '6. Save output/output_1702835621.mp4',
        '7. QMediaPlayer plays video 🎥'
    ]
    for step in workflow2:
        doc.add_paragraph(step)
    
    # Technical Features
    doc.add_page_break()
    doc.add_heading('🎯 Technical Features', 1)
    
    doc.add_heading('1. Multi-threaded Architecture', 2)
    threads = [
        'CameraWorker: Camera processing',
        'PredictionWorker: AI prediction',
        'TTSWorker: Speech synthesis',
        'Main thread: GUI rendering'
    ]
    for thread in threads:
        doc.add_paragraph(thread, style='List Bullet')
    
    doc.add_heading('2. AI Integration', 2)
    doc.add_paragraph('Local Ollama (Qwen2.5:7b) with 4 AI functions:')
    ai_functions = [
        'Word Prediction: Next word or auto-complete',
        'Grammar Correction: Fix recognition errors',
        'ASL Translation: English → ASL Gloss',
        'Similarity Search: Find alternative words'
    ]
    for func in ai_functions:
        doc.add_paragraph(func, style='List Bullet')
    
    doc.add_heading('3. Error Handling', 2)
    error_handling = [
        'API timeout protection (5-15 seconds)',
        'Missing video fallback to similar words',
        'AI failure downgrade to original text',
        'Thread-safe signal-slot communication'
    ]
    for eh in error_handling:
        doc.add_paragraph(eh, style='List Bullet')
    
    doc.add_heading('4. User Experience', 2)
    ux_features = [
        'Real-time preview (progress bar shows 0.6s hold)',
        'Cursor-based editing (arrow keys to move)',
        'Keyboard shortcuts (F1-F3 for suggestions)',
        'Dark theme interface',
        'Status indicators for all operations'
    ]
    for ux in ux_features:
        doc.add_paragraph(ux, style='List Bullet')
    
    # Project Statistics
    doc.add_heading('📊 Project Statistics', 1)
    stats = [
        ('Total Lines of Code', '~1000+ lines'),
        ('Supported Gestures', '26 letters (A-Z)'),
        ('Video Database', '2000+ sign language videos'),
        ('AI Models', '3 models (KeyPointClassifier, Qwen2.5, InsightFace)'),
        ('Supported Features', 'Recognition, Translation, Generation, Speech, Face Swap')
    ]
    
    table = doc.add_table(rows=len(stats) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    for i, (metric, value) in enumerate(stats):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    # AI Technology Highlights
    doc.add_page_break()
    doc.add_heading('🧠 Ollama AI Technology Highlights', 1)
    
    doc.add_heading('1. Local LLM Deployment', 2)
    doc.add_paragraph('Model: Qwen2.5:7b (7 billion parameters)', style='List Bullet')
    doc.add_paragraph('Deployment: Local (localhost:11434)', style='List Bullet')
    doc.add_paragraph('Advantages:', style='Heading 5')
    advantages = [
        '✅ No internet required - works offline',
        '✅ Data privacy protection - sensitive info stays local',
        '✅ No API costs - completely free',
        '✅ Low latency response (<5 seconds)'
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_heading('2. Four AI Function Modules', 2)
    
    doc.add_paragraph('A. Smart Word Prediction 📝', style='Heading 4')
    doc.add_paragraph('Predicts next word or completes current word based on context')
    doc.add_paragraph('Features:', style='List Bullet')
    doc.add_paragraph('Real-time prediction (0.3s refresh)', style='List Bullet 2')
    doc.add_paragraph('Low temperature (0.2) ensures accuracy', style='List Bullet 2')
    doc.add_paragraph('Returns top 3 suggestions', style='List Bullet 2')
    
    doc.add_paragraph('B. Grammar Auto-correction ✍️', style='Heading 4')
    doc.add_paragraph('Fixes fragmented text from sign recognition into complete sentences')
    doc.add_paragraph('Features:', style='List Bullet')
    doc.add_paragraph('15-second timeout for response', style='List Bullet 2')
    doc.add_paragraph('Fallback to original text on failure', style='List Bullet 2')
    
    doc.add_paragraph('C. ASL Gloss Translation 🤟', style='Heading 4')
    doc.add_paragraph('Converts English to ASL sign language grammar')
    doc.add_paragraph('Features:', style='List Bullet')
    doc.add_paragraph('Word order restructuring (English → ASL)', style='List Bullet 2')
    doc.add_paragraph('Removes redundant words (the, is, are)', style='List Bullet 2')
    doc.add_paragraph('Matches with sign video database', style='List Bullet 2')
    
    doc.add_paragraph('D. Semantic Similarity Search 🔍', style='Heading 4')
    doc.add_paragraph('Finds semantically similar words when database is missing a word')
    doc.add_paragraph('Example: Missing "hi" → Uses "hello" instead', style='List Bullet')
    doc.add_paragraph('Ultra-low temperature (0.1) ensures semantic accuracy', style='List Bullet')
    
    doc.add_heading('3. Comparison with Traditional Methods', 2)
    comparison = doc.add_table(rows=5, cols=3)
    comparison.style = 'Light Grid Accent 1'
    
    comp_data = [
        ['Function', 'Traditional Method', 'Ollama AI Method'],
        ['Word Prediction', 'Fixed dictionary', 'Context understanding'],
        ['Grammar Correction', 'None', 'AI auto-correction'],
        ['Sign Translation', 'Word-by-word', 'Grammar restructuring'],
        ['Similar Word Search', 'String matching', 'Semantic understanding']
    ]
    
    for i, row_data in enumerate(comp_data):
        row_cells = comparison.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Conclusion
    doc.add_page_break()
    doc.add_heading('✨ Conclusion', 1)
    doc.add_paragraph(
        'This is a complete AAC (Augmentative and Alternative Communication) system '
        'designed for the hearing-impaired community. It combines cutting-edge AI technology '
        'with real-time computer vision to bridge the communication gap between sign language '
        'users and the hearing world. The system demonstrates:'
    )
    
    conclusion_points = [
        'Practical application of local LLM (Ollama)',
        'Real-time gesture recognition using MediaPipe and TensorFlow',
        'Seamless integration of multiple AI models',
        'User-friendly interface with PyQt6',
        'Robust error handling and fallback mechanisms',
        'Privacy-focused offline operation'
    ]
    for point in conclusion_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'Testing_Folder_Documentation.docx')
    doc.save(output_path)
    print(f"✅ Documentation created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    try:
        file_path = create_documentation()
        print(f"\n📄 Document saved at: {file_path}")
        print("You can now open this file in Microsoft Word or convert it to PDF.")
    except ImportError:
        print("❌ Error: python-docx not installed")
        print("Please run: pip install python-docx")
    except Exception as e:
        print(f"❌ Error creating documentation: {e}")
